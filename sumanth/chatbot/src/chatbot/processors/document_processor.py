"""
Document processor for handling multiple file formats (PDF, DOCX, TXT, MD).
Supports semantic chunking and metadata extraction for better RAG performance.
"""
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from langchain.docstore.document import Document

from chatbot.utils.logger import logger
from chatbot.utils.exceptions import EmbeddingError


class DocumentProcessor:
    """Process documents from various formats and extract text with metadata."""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md', '.markdown'}
    
    def __init__(self):
        """Initialize the document processor."""
        self._load_extractors()
    
    def _load_extractors(self):
        """Lazy load extractors only when needed."""
        self._pdf_extractor = None
        self._docx_extractor = None
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            import PyPDF2
        except ImportError:
            raise EmbeddingError(
                "PyPDF2 is required for PDF processing. "
                "Install it with: pip install PyPDF2"
            )
        
        text = ""
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += f"\n\n--- Page {page_num + 1} ---\n\n{page_text}"
        except Exception as e:
            logger.error(f"Error extracting PDF {file_path}: {e}")
            raise EmbeddingError(f"Failed to extract text from PDF: {e}")
        
        return text.strip()
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise EmbeddingError(
                "python-docx is required for DOCX processing. "
                "Install it with: pip install python-docx"
            )
        
        try:
            doc = DocxDocument(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Error extracting DOCX {file_path}: {e}")
            raise EmbeddingError(f"Failed to extract text from DOCX: {e}")
    
    def _extract_txt(self, file_path: Path) -> str:
        """Extract text from plain text file."""
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Fallback to latin-1
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
            except Exception as e:
                logger.error(f"Error extracting TXT {file_path}: {e}")
                raise EmbeddingError(f"Failed to extract text from TXT: {e}")
    
    def _extract_markdown(self, file_path: Path) -> str:
        """Extract text from Markdown file."""
        return self._extract_txt(file_path)  # Same as TXT for now
    
    def extract_text(self, file_path: Path) -> str:
        """
        Extract text from a document file.
        
        Args:
            file_path: Path to the document file
        
        Returns:
            Extracted text content
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise EmbeddingError(f"File not found: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return self._extract_pdf(file_path)
        elif suffix == '.docx':
            return self._extract_docx(file_path)
        elif suffix in {'.txt', '.text'}:
            return self._extract_txt(file_path)
        elif suffix in {'.md', '.markdown'}:
            return self._extract_markdown(file_path)
        else:
            raise EmbeddingError(
                f"Unsupported file format: {suffix}. "
                f"Supported formats: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )
    
    def process_documents(
        self,
        input_path: str,
        recursive: bool = False
    ) -> List[Document]:
        """
        Process one or more documents and return Document objects.
        
        Args:
            input_path: Path to file or directory
            recursive: If True, process files recursively in directories
        
        Returns:
            List of Document objects with metadata
        """
        input_path = Path(input_path)
        documents = []
        
        if input_path.is_file():
            # Single file
            files = [input_path]
        elif input_path.is_dir():
            # Directory - find all supported files
            pattern = "**/*" if recursive else "*"
            files = []
            for ext in self.SUPPORTED_EXTENSIONS:
                files.extend(input_path.glob(f"{pattern}{ext}"))
        else:
            raise EmbeddingError(f"Path does not exist: {input_path}")
        
        if not files:
            raise EmbeddingError(f"No supported documents found in: {input_path}")
        
        logger.info(f"Processing {len(files)} document(s)...")
        
        for file_path in files:
            try:
                logger.info(f"Processing: {file_path.name}")
                text = self.extract_text(file_path)
                
                if not text.strip():
                    logger.warning(f"No text extracted from {file_path.name}")
                    continue
                
                # Create document with metadata
                doc = Document(
                    page_content=text,
                    metadata={
                        'source': str(file_path),
                        'filename': file_path.name,
                        'file_type': file_path.suffix.lower(),
                        'file_size': file_path.stat().st_size,
                    }
                )
                documents.append(doc)
                logger.info(f"Extracted {len(text)} characters from {file_path.name}")
                
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
                continue
        
        if not documents:
            raise EmbeddingError("No documents were successfully processed")
        
        logger.info(f"Successfully processed {len(documents)} document(s)")
        return documents


class SemanticChunker:
    """
    Advanced chunking that preserves semantic meaning.
    Uses sentence boundaries and paragraph breaks for better context.
    """
    
    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        min_chunk_size: int = 50,
        separators: Optional[List[str]] = None
    ):
        """
        Initialize semantic chunker.
        
        Args:
            chunk_size: Target chunk size in words
            chunk_overlap: Overlap between chunks in words
            min_chunk_size: Minimum chunk size to keep
            separators: List of separators to use (in order of preference)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
        
        # Default separators: try to break at semantic boundaries
        self.separators = separators or [
            "\n\n\n",  # Multiple newlines (section breaks)
            "\n\n",    # Paragraph breaks
            ".\n",     # Sentence + newline
            ". ",      # Sentence endings
            "! ",      # Exclamation
            "? ",      # Question
            "\n",      # Single newline
            " ",       # Space
            "",        # Character level (last resort)
        ]
    
    def _word_count(self, text: str) -> int:
        """Count words in text."""
        return len(text.split())
    
    def _split_text(self, text: str, separator: str) -> List[str]:
        """Split text by separator."""
        if separator == "":
            return list(text)
        return text.split(separator)
    
    def chunk_text(self, text: str) -> List[str]:
        """
        Split text into semantically meaningful chunks.
        
        Args:
            text: Text to chunk
        
        Returns:
            List of text chunks
        """
        if not text.strip():
            return []
        
        # Try each separator in order
        for separator in self.separators:
            chunks = self._split_text(text, separator)
            
            # Check if chunks are appropriate size
            if len(chunks) > 1:
                # Check if we can use this separator
                max_chunk_size = max(self._word_count(chunk) for chunk in chunks if chunk.strip())
                
                if max_chunk_size <= self.chunk_size * 1.5:  # Allow some flexibility
                    # Merge chunks that are too small
                    merged_chunks = self._merge_chunks(chunks, separator)
                    return merged_chunks
        
        # If no separator worked well, fall back to character-level splitting
        return self._split_by_size(text)
    
    def _merge_chunks(self, chunks: List[str], separator: str) -> List[str]:
        """Merge small chunks together."""
        merged = []
        current_chunk = ""
        current_size = 0
        
        for chunk in chunks:
            chunk = chunk.strip()
            if not chunk:
                continue
            
            chunk_size = self._word_count(chunk)
            
            if current_size + chunk_size <= self.chunk_size:
                # Add to current chunk
                if current_chunk:
                    current_chunk += separator + chunk
                else:
                    current_chunk = chunk
                current_size += chunk_size
            else:
                # Save current chunk and start new one
                if current_chunk and current_size >= self.min_chunk_size:
                    merged.append(current_chunk)
                
                # Handle oversized chunks
                if chunk_size > self.chunk_size:
                    # Split oversized chunk
                    sub_chunks = self._split_oversized_chunk(chunk)
                    merged.extend(sub_chunks[:-1])  # Add all but last
                    current_chunk = sub_chunks[-1]  # Start new chunk with last
                    current_size = self._word_count(current_chunk)
                else:
                    current_chunk = chunk
                    current_size = chunk_size
        
        # Add final chunk
        if current_chunk and current_size >= self.min_chunk_size:
            merged.append(current_chunk)
        
        # Add overlap between chunks
        return self._add_overlap(merged)
    
    def _split_oversized_chunk(self, chunk: str) -> List[str]:
        """Split a chunk that's too large."""
        words = chunk.split()
        chunks = []
        
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunks.append(" ".join(chunk_words))
        
        return chunks
    
    def _split_by_size(self, text: str) -> List[str]:
        """Fallback: split by size when semantic splitting fails."""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunks.append(" ".join(chunk_words))
        
        return chunks
    
    def _add_overlap(self, chunks: List[str]) -> List[str]:
        """Add overlap between chunks for better context."""
        if len(chunks) <= 1 or self.chunk_overlap == 0:
            return chunks
        
        overlapped = [chunks[0]]  # First chunk as-is
        
        for i in range(1, len(chunks)):
            prev_chunk = chunks[i - 1]
            current_chunk = chunks[i]
            
            # Get last N words from previous chunk
            prev_words = prev_chunk.split()
            overlap_words = prev_words[-self.chunk_overlap:] if len(prev_words) > self.chunk_overlap else prev_words
            
            # Prepend overlap to current chunk
            overlapped_chunk = " ".join(overlap_words) + " " + current_chunk
            overlapped.append(overlapped_chunk)
        
        return overlapped
    
    def chunk_documents(self, documents: List[Document]) -> List[Document]:
        """
        Chunk a list of documents.
        
        Args:
            documents: List of Document objects
        
        Returns:
            List of chunked Document objects with preserved metadata
        """
        chunked_docs = []
        
        for doc in documents:
            chunks = self.chunk_text(doc.page_content)
            
            for i, chunk in enumerate(chunks):
                # Preserve metadata and add chunk info
                chunk_metadata = doc.metadata.copy()
                chunk_metadata['chunk_index'] = i
                chunk_metadata['total_chunks'] = len(chunks)
                
                chunked_doc = Document(
                    page_content=chunk,
                    metadata=chunk_metadata
                )
                chunked_docs.append(chunked_doc)
        
        return chunked_docs

